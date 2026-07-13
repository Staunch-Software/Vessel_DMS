"""Text extraction for Word/Excel documents.

Kept separate from `extract.py` since these deps (python-docx, openpyxl) are
lighter than the PaddleOCR/PyMuPDF stack but still worth lazy-importing so a
stub-mode/dev install without them doesn't fail at import time.
"""
from __future__ import annotations

import io


def docx_text(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def xlsx_text(file_bytes: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(sheet.title)
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)
